"""
schablone_service.py
--------------------
Generates an Aufnahme-Schablone .docx file and saves it to outputs/.

Strategy:
- If template/aufnahme_template.docx exists, copy it and return the copy.
- Otherwise generate a fallback document that includes all required section
  markers so that document_service.insert_blocks_into_section() can find them.

Required markers (searched by startswith):
  "Derzeitige Beschwerden" — start of the somatik section
  "Vormedikation"          — end marker (exclusive)
"""

from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path


class SchabloneService:

    _OUTPUT_DIR  = Path(__file__).parent.parent / "outputs"
    _TEMPLATE    = Path(__file__).parent.parent / "templates" / "aufnahme_template.docx"

    def generate_and_save(self) -> Path:
        self._OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename  = f"Aufnahme_Schablone_{timestamp}.docx"
        path      = self._OUTPUT_DIR / filename

        if self._TEMPLATE.exists():
            shutil.copy2(self._TEMPLATE, path)
            return path

        return self._generate_fallback(path)

    def _generate_fallback(self, path: Path) -> Path:
        from docx import Document

        doc = Document()
        doc.add_heading("Aufnahme TCM Klinik", level=1)
        doc.add_heading("Derzeitige Beschwerden (somatisch)", level=2)
        doc.add_paragraph("")
        doc.add_heading("Vormedikation", level=2)
        doc.add_paragraph("")
        doc.add_heading("Psychosomatische Anamnese", level=2)
        doc.add_paragraph("")
        doc.add_heading("Diagnosen", level=2)
        doc.add_paragraph("")
        doc.save(path)

        return path
